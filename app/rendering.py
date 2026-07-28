from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from .domain import format_timestamp
from .minutes_templates import (
    default_minutes_template,
    normalize_minutes_template,
)


def render_transcript_markdown(
    meeting: dict[str, Any],
    segments: list[dict[str, Any]],
    speakers: dict[str, str],
) -> str:
    lines = [
        f"# {meeting['title']}：逐字稿",
        "",
        f"- 日期：{meeting['meeting_date']}",
        f"- 预计发言人数：{meeting['expected_speakers']}",
        "",
    ]
    for segment in segments:
        speaker = speakers.get(segment["speaker"]) or segment["speaker"]
        timestamp = segment.get("timestamp") or format_timestamp(
            segment.get("start")
        )
        lines.extend(
            [f"## [{timestamp}] {speaker}", "", segment["text"].strip(), ""]
        )
    return "\n".join(lines).rstrip() + "\n"


def render_transcript_text(
    meeting: dict[str, Any],
    segments: list[dict[str, Any]],
    speakers: dict[str, str],
) -> str:
    lines = [
        f"{meeting['title']}：录音转写",
        f"日期：{meeting['meeting_date']}",
        f"预计发言人数：{meeting['expected_speakers']}",
        "",
    ]
    for segment in segments:
        speaker = speakers.get(segment["speaker"]) or segment["speaker"]
        timestamp = segment.get("timestamp") or format_timestamp(
            segment.get("start")
        )
        lines.extend(
            [
                f"[{timestamp}] {speaker}",
                str(segment.get("text") or "").strip(),
                "",
            ]
        )
    return "\n".join(lines).rstrip() + "\n"


def render_minutes_markdown(minutes: dict[str, Any]) -> str:
    meeting = minutes.get("meeting", {})
    lines = [
        f"# {meeting.get('title', '会议纪要')}",
        "",
        f"- 日期：{meeting.get('date', '未明确')}",
        "",
    ]
    for section in minutes_sections(minutes):
        key = section["key"]
        title = section["title"]
        lines.extend([f"## {title}", ""])
        if section["kind"] == "summary":
            lines.extend(
                [str(minutes.get(key) or "未明确"), ""]
            )
        elif section["kind"] == "actions":
            _append_actions_markdown(lines, minutes.get(key, []))
        else:
            _append_list(lines, minutes.get(key, []))
    return "\n".join(lines).rstrip() + "\n"


def _append_actions_markdown(lines: list[str], actions: Any) -> None:
    if not isinstance(actions, list) or not actions:
        lines.extend(["- 无明确待办事项", ""])
        return
    lines.extend(
        [
            "| 负责人 | 任务 | 截止日期 | 原文时间 | 状态 |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for item in actions:
        if not isinstance(item, dict):
            item = {"task": str(item)}
        lines.append(
            "| {owner} | {task} | {due} | {evidence_time} | {status} |".format(
                **{
                    key: _escape_table(str(item.get(key, "未明确")))
                    for key in (
                        "owner",
                        "task",
                        "due",
                        "evidence_time",
                        "status",
                    )
                }
            )
        )
    lines.append("")


def render_minutes_text(minutes: dict[str, Any]) -> str:
    markdown = render_minutes_markdown(minutes)
    text = re.sub(r"^#{1,6}\s+", "", markdown, flags=re.MULTILINE)
    text = text.replace("**", "").replace("`", "")
    return text


def render_minutes_docx(minutes: dict[str, Any], destination: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    meeting = minutes.get("meeting", {})
    document.add_heading(str(meeting.get("title") or "会议纪要"), level=0)
    document.add_paragraph(f"日期：{meeting.get('date', '未明确')}")
    for section in minutes_sections(minutes):
        key = section["key"]
        title = section["title"]
        document.add_heading(title, level=1)
        if section["kind"] == "summary":
            document.add_paragraph(
                str(minutes.get(key) or "未明确")
            )
        elif section["kind"] == "actions":
            _add_docx_actions(document, minutes.get(key, []))
        else:
            _add_docx_list(document, minutes.get(key, []))
    document.save(destination)


def _add_docx_actions(document: Document, actions: Any) -> None:
    if not isinstance(actions, list) or not actions:
        document.add_paragraph("无明确待办事项")
        return
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("负责人", "任务", "截止日期", "原文时间", "状态")
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for item in actions:
        if not isinstance(item, dict):
            item = {"task": str(item)}
        cells = table.add_row().cells
        values = (
            item.get("owner", "未明确"),
            item.get("task", "待确认"),
            item.get("due", "未明确"),
            item.get("evidence_time", "未明确"),
            item.get("status", "待处理"),
        )
        for cell, value in zip(cells, values):
            cell.text = str(value)


def _append_list(lines: list[str], items: Any) -> None:
    if not isinstance(items, list) or not items:
        lines.extend(["- 未明确", ""])
        return
    for item in items:
        lines.append(f"- {_item_to_text(item)}")
    lines.append("")


def _add_docx_list(document: Document, items: Any) -> None:
    if not isinstance(items, list) or not items:
        document.add_paragraph("未明确")
        return
    for item in items:
        document.add_paragraph(_item_to_text(item), style="List Bullet")


def _item_to_text(item: Any) -> str:
    if not isinstance(item, dict):
        return str(item)
    evidence = item.get("evidence_time")
    main = minutes_item_text(item)
    return f"{main}（原文时间：{evidence}）" if evidence else main


def minutes_item_text(item: Any) -> str:
    if not isinstance(item, dict):
        return str(item)
    preferred = (
        "content",
        "progress",
        "result",
        "suggestion",
        "decision",
        "question",
        "followup",
        "text",
    )
    main = next(
        (str(item[key]) for key in preferred if item.get(key)),
        "；".join(
            f"{key}: {value}"
            for key, value in item.items()
            if key != "evidence_time"
        ),
    )
    return main


def minutes_sections(minutes: dict[str, Any]) -> list[dict[str, str]]:
    raw_template = minutes.get("template")
    try:
        template = normalize_minutes_template(
            raw_template
            if isinstance(raw_template, dict)
            else default_minutes_template()
        )
    except (TypeError, ValueError):
        template = default_minutes_template()
    return list(template["sections"])


def _escape_table(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")
